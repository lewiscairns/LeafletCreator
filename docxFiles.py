# Import necessary libraries and modules
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import tkinter as tk


# This function is used to create the DOCX document
def create_document(title, all_rows, folder_selected, font_style, font_size, watermark_image, watermark_text):
    # Create the document
    document = Document()

    # Set the document style
    style = document.styles['Normal']
    font = style.font
    font.name = font_style
    font.size = Pt(font_size)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5

    # Set the main body, header, and footer
    section = document.sections[0]
    header = section.header
    footer = section.footer

    paragraph = header.paragraphs[0]
    paragraph.alignment = 1
    run = paragraph.add_run()
    run.add_text(title)
    paragraph = footer.paragraphs[0]
    paragraph.alignment = 1
    run = paragraph.add_run()
    run.add_text(watermark_text)
    run.add_text("   ")
    run.add_picture(watermark_image, width=Inches(0.20), height=Inches(0.20))

    # Add the images and text from each page to the document
    for page in all_rows:
        # Add the document's table
        table = document.add_table(rows=0, cols=2)
        # For each row in the page, add the image and text to the table
        for row in page[1:]:
            bold_word_positions = row[1].tag_ranges("bold")
            row_cells = table.add_row().cells
            image = row_cells[0].paragraphs[0]
            text = row_cells[1].paragraphs[0]
            run = image.add_run()
            run.add_picture(row[0], width=Inches(1.80), height=Inches(1.80))
            sentence = retrieve_input(row[1])
            # Add the text to the document, checking if it should be made bold
            for word in sentence:
                word, position = word.split("#")
                if str(position) in str(bold_word_positions):
                    text.add_run(str(word)).bold = True
                else:
                    text.add_run(str(word))
                text.add_run(" ")

    # Save the document
    document.save(folder_selected + "/" + title + '.docx')


# This function is used to retrieve the text from the text box
def retrieve_input(text_box):
    words = []
    start_index = "1.0"
    text_box.insert(tk.END, " ")
    # Get the position of each word in the text box
    while True:
        # Get the position of the next space
        space_index = text_box.search(" ", start_index, tk.END)

        # If there are no more spaces, break the loop
        if space_index == "":
            break

        # Get the word and its position
        word_position = space_index

        # Get the line and column of the word
        line, column = word_position.split(".")
        position = f"{line}.{column}"

        # Get the word using the positions
        word = text_box.get(start_index, space_index)

        # Add the word and its position to the list, splitting it using #
        words.append(word + "#" + position)
        start_index = word_position + "+1c"

    # Return the list of words and their positions
    return words
